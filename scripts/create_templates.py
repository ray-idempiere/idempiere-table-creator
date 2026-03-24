#!/usr/bin/env python3
"""Generate iDempiere table creator asset templates."""
import os
import openpyxl
from openpyxl.styles import Font, PatternFill

ASSETS_DIR = os.path.expanduser("~/.claude/skills/idempiere-table-creator/assets")


def make_excel_template():
    wb = openpyxl.Workbook()

    # ── Sheet 1: Table Info ──────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Table Info"

    header_font = Font(bold=True)
    header_fill = PatternFill("solid", fgColor="DDEBF7")
    input_fill  = PatternFill("solid", fgColor="FFFFE0")

    rows = [
        ("Field",           "Value",                  "Notes"),
        ("Prefix",          "",                       "Module prefix without underscore, e.g. HR, CVG"),
        ("Table Name",      "",                       "PascalCase, e.g. EmployeeCard (prefix added automatically)"),
        ("Description",     "",                       "Short description for SQL comment"),
        ("Entity Type",     "U",                      "U = user-defined; override for specific module"),
        ("Is Doc Enabled",  "N",                      "Y = add DocStatus/DocAction workflow columns"),
    ]

    ws1.column_dimensions["A"].width = 20
    ws1.column_dimensions["B"].width = 35
    ws1.column_dimensions["C"].width = 55

    for i, (field, value, note) in enumerate(rows, start=1):
        ws1.cell(i, 1, field).font = header_font
        if i == 1:
            ws1.cell(i, 1).fill = header_fill
            ws1.cell(i, 2).fill = header_fill
            ws1.cell(i, 3).fill = header_fill
        else:
            ws1.cell(i, 2, value).fill = input_fill

        ws1.cell(i, 3, note).font = Font(italic=True, color="666666")

    # ── Sheet 2: Columns ────────────────────────────────────────
    ws2 = wb.create_sheet("Columns")

    col_headers = [
        "ColumnName", "DataType", "Length", "Mandatory",
        "FK Table", "Chinese Name (zh_TW)", "Description / Help",
    ]
    col_widths = [25, 15, 10, 12, 25, 25, 40]

    for j, (h, w) in enumerate(zip(col_headers, col_widths), start=1):
        cell = ws2.cell(1, j, h)
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DDEBF7")
        ws2.column_dimensions[chr(64 + j)].width = w

    examples = [
        ("EmployeeNo",    "String",  "30", "N", "",             "員工編號",   "Unique employee identifier"),
        ("Name",          "String",  "60", "Y", "",             "姓名",       "Full name"),
        ("BirthDate",     "Date",    "",   "N", "",             "生日",       ""),
        ("Gender",        "List",    "1",  "N", "",             "性別",       "M=Male, F=Female"),
        ("C_BPartner_ID", "ID",      "",   "N", "C_BPartner",   "業務夥伴",   ""),
    ]
    for r, row in enumerate(examples, start=2):
        for c, val in enumerate(row, start=1):
            ws2.cell(r, c, val).fill = PatternFill("solid", fgColor="FFFFE0")

    # ── Sheet 3: Instructions ────────────────────────────────────
    ws3 = wb.create_sheet("Instructions")
    instructions = [
        "iDempiere Table Creator — Template Instructions",
        "",
        "Sheet 1 (Table Info):",
        "  Prefix        — module prefix WITHOUT underscore, e.g. HR, CVG, IVG",
        "  Table Name    — PascalCase name WITHOUT prefix, e.g. EmployeeCard",
        "  Is Doc Enabled — Y to add DocStatus/DocAction workflow columns",
        "",
        "Sheet 2 (Columns):",
        "  ColumnName    — PascalCase, max 30 chars, no spaces",
        "  DataType      — String / Date / Amount / ID / YesNo / List / Number / Quantity",
        "  Length        — required for String type; leave blank for others",
        "  Mandatory     — Y or N",
        "  FK Table      — if DataType=ID, the referenced table name",
        "  Chinese Name  — zh_TW translation for AD_Element_Trl (IsTranslated=Y if filled)",
        "",
        "Do NOT add mandatory columns: AD_Client_ID, AD_Org_ID, IsActive, Created,",
        "CreatedBy, Updated, UpdatedBy — these are auto-generated.",
    ]
    ws3.column_dimensions["A"].width = 80
    for r, text in enumerate(instructions, start=1):
        cell = ws3.cell(r, 1, text)
        if r == 1:
            cell.font = Font(bold=True, size=13)

    path = os.path.join(ASSETS_DIR, "template.xlsx")
    wb.save(path)
    print(f"Created: {path}")


def make_word_template():
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading("iDempiere Table Specification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Fill in this document and run /create-table <path-to-this-file>")
    doc.add_paragraph("")

    # Section 1: Table Metadata
    doc.add_heading("1. Table Metadata", level=1)

    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Table Grid"

    meta_rows = [
        ("Field",         "Value"),
        ("Prefix",        "(e.g. HR, CVG — no underscore)"),
        ("Table Name",    "(PascalCase, no prefix, e.g. EmployeeCard)"),
        ("Description",   ""),
        ("Entity Type",   "U"),
        ("Is Doc Enabled","N  (Y = add DocStatus/DocAction workflow columns)"),
    ]
    for i, (field, hint) in enumerate(meta_rows):
        meta_table.rows[i].cells[0].text = field
        meta_table.rows[i].cells[1].text = hint
        if i == 0:
            for cell in meta_table.rows[i].cells:
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True

    doc.add_paragraph("")

    # Section 2: Column Definitions
    doc.add_heading("2. Column Definitions", level=1)
    doc.add_paragraph(
        "List each column below. Do NOT include mandatory columns "
        "(AD_Client_ID, AD_Org_ID, IsActive, Created/By, Updated/By) — "
        "they are auto-generated."
    )

    col_table = doc.add_table(rows=6, cols=7)
    col_table.style = "Table Grid"

    col_headers = [
        "ColumnName", "DataType", "Length", "Mandatory",
        "FK Table", "Chinese Name\n(zh_TW)", "Description"
    ]
    header_row = col_table.rows[0]
    for j, h in enumerate(col_headers):
        cell = header_row.cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)

    examples = [
        ("EmployeeNo",   "String", "30", "N", "",           "員工編號", ""),
        ("Name",         "String", "60", "Y", "",           "姓名",     ""),
        ("BirthDate",    "Date",   "",   "N", "",           "生日",     ""),
        ("Gender",       "List",   "1",  "N", "",           "性別",     "M=Male, F=Female"),
        ("C_BPartner_ID","ID",     "",   "N", "C_BPartner", "業務夥伴", ""),
    ]
    for r, row_data in enumerate(examples, start=1):
        for c, val in enumerate(row_data):
            cell = col_table.rows[r].cells[c]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_paragraph("")

    # Section 3: Repeating Sections
    doc.add_heading("3. Repeating Sections (Child Tables)", level=1)
    doc.add_paragraph(
        "If this table has repeating sections (e.g., education history, work experience), "
        "list them here. Each section becomes a separate child table."
    )

    rep_table = doc.add_table(rows=4, cols=2)
    rep_table.style = "Table Grid"
    rep_table.rows[0].cells[0].text = "Section Name (Chinese)"
    rep_table.rows[0].cells[1].text = "English Table Suffix"
    for cell in rep_table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    path = os.path.join(ASSETS_DIR, "template.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    os.makedirs(ASSETS_DIR, exist_ok=True)
    make_excel_template()
    make_word_template()
